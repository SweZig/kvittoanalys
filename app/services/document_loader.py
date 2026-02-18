"""Document loader — converts uploaded files into images for analysis."""

import base64
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from PIL import Image

from app.config import settings


class DocumentLoader:
    """Load and convert documents/images into a format ready for AI analysis."""

    @staticmethod
    def load_file(file_path: Path) -> list[dict]:
        """
        Load a file and return a list of content blocks.

        Each block is a dict with:
            - type: "image" | "text"
            - data: base64 string (for images) or plain text
            - media_type: MIME type (for images)
            - page: page number (1-indexed)
            - source: original filename
        """
        suffix = file_path.suffix.lower()

        if suffix in settings.supported_image_types:
            return DocumentLoader._load_image(file_path)
        elif suffix == ".pdf":
            return DocumentLoader._load_pdf(file_path)
        elif suffix in (".docx", ".doc"):
            return DocumentLoader._load_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    @staticmethod
    def _load_image(file_path: Path) -> list[dict]:
        """Load a single image file."""
        media_type_map = {
            ".png": "image/png",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".gif": "image/gif",
            ".webp": "image/webp",
            ".bmp": "image/png",  # convert to PNG
            ".tiff": "image/png",  # convert to PNG
        }

        suffix = file_path.suffix.lower()
        media_type = media_type_map.get(suffix, "image/png")

        # Convert non-standard formats to PNG
        if suffix in (".bmp", ".tiff"):
            img = Image.open(file_path).convert("RGB")
            import io

            buffer = io.BytesIO()
            img.save(buffer, format="PNG")
            image_data = base64.standard_b64encode(buffer.getvalue()).decode("utf-8")
        else:
            image_data = base64.standard_b64encode(file_path.read_bytes()).decode(
                "utf-8"
            )

        return [
            {
                "type": "image",
                "data": image_data,
                "media_type": media_type,
                "page": 1,
                "source": file_path.name,
            }
        ]

    @staticmethod
    def _load_pdf(file_path: Path, dpi: int = 200) -> list[dict]:
        """Convert each page of a PDF to an image."""
        blocks = []
        doc = fitz.open(file_path)

        for page_num in range(len(doc)):
            page = doc[page_num]
            # Render page to image
            zoom = dpi / 72
            matrix = fitz.Matrix(zoom, zoom)
            pixmap = page.get_pixmap(matrix=matrix)
            image_data = base64.standard_b64encode(pixmap.tobytes("png")).decode(
                "utf-8"
            )

            blocks.append(
                {
                    "type": "image",
                    "data": image_data,
                    "media_type": "image/png",
                    "page": page_num + 1,
                    "source": file_path.name,
                }
            )

        doc.close()
        return blocks

    @staticmethod
    def _load_docx(file_path: Path) -> list[dict]:
        """Extract text content from a Word document."""
        doc = DocxDocument(file_path)
        full_text = []

        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))

        return [
            {
                "type": "text",
                "data": "\n\n".join(full_text),
                "page": 1,
                "source": file_path.name,
            }
        ]
